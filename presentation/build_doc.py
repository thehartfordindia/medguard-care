"""
Generate a polished, client-ready MedGuard AI Word document.

Run:
    python presentation/build_doc.py

Output:
    presentation/MedGuard_AI.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Hartford brand palette ───────────────────────────────
HARTFORD_RED  = RGBColor(0xE4, 0x00, 0x2B)
NAVY          = RGBColor(0x0A, 0x25, 0x40)
INK           = RGBColor(0x1F, 0x29, 0x37)
GRAY          = RGBColor(0x6B, 0x72, 0x80)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
TEAL          = RGBColor(0x14, 0xB8, 0xA6)


# ─── XML helpers ──────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def page_border(section, color="E4002B", size=24):
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "24")
        b.set(qn("w:color"), color)
        pg_borders.append(b)
    sect_pr.append(pg_borders)


def add_horizontal_rule(doc, color="E4002B", size=12):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


# ─── Styled writers ───────────────────────────────────────
def add_run(para, text, *, size=11, bold=False, color=INK, font="Calibri"):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    return run


def add_heading(doc, text, *, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    if level == 0:
        add_run(p, text, size=28, bold=True, color=NAVY)
    elif level == 1:
        add_run(p, text, size=18, bold=True, color=HARTFORD_RED)
    elif level == 2:
        add_run(p, text, size=14, bold=True, color=NAVY)
    else:
        add_run(p, text, size=12, bold=True, color=INK)
    return p


def add_body(doc, text, *, size=11, bold=False, color=INK, align=None,
             space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    add_run(p, text, size=size, bold=bold, color=color)
    return p


def add_bullets(doc, items, *, size=11, color=INK, indent=Inches(0.25)):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = indent
        p.paragraph_format.space_after = Pt(2)
        add_run(p, item, size=size, color=color)


def add_table(doc, headers, rows, *, header_fill="0A2540",
               alt_fill="F3F4F6", first_col_bold=False):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.autofit = True
    tbl.style = "Light Grid Accent 1"

    # Header row
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        cell.text = ""
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, h, size=11, bold=True, color=WHITE)

    # Body rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = tbl.rows[r + 1].cells[c]
            cell.text = ""
            if r % 2 == 1:
                shade_cell(cell, alt_fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, val, size=10,
                    bold=(c == 0 and first_col_bold),
                    color=NAVY if (c == 0 and first_col_bold) else INK)
    return tbl


def add_callout_box(doc, title, body, *, fill="0A2540", title_color=WHITE,
                     body_color=WHITE):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    shade_cell(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Title
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    add_run(p1, title, size=11, bold=True,
            color=RGBColor(0xFF, 0xC1, 0xC9) if fill == "0A2540" else title_color)
    # Body
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, body, size=14, bold=True, color=body_color)
    add_body(doc, "", size=2)


# ═════════════════════════════════════════════════════════════
# DOCUMENT BUILDER
# ═════════════════════════════════════════════════════════════
def build():
    doc = Document()

    # Page setup — letter, narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        page_border(section, color="E4002B", size=18)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    # ── COVER ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "THE HARTFORD  ·  AI / ML / DATA SCIENCE",
            size=10, bold=True, color=HARTFORD_RED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "MedGuard AI", size=44, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Intelligent Claim Intelligence — powered by AI-Augmented Risk Mitigation",
            size=15, color=GRAY)

    add_horizontal_rule(doc, color="E4002B", size=18)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "A reusable AI pattern for Middle & Large Commercial and Group Benefits",
            size=12, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Presented by  ·  [Your Name]    |    Hackathon 2026  ·  Internal Demo",
            size=10, color=GRAY)

    # ── EXECUTIVE SUMMARY ─────────────────────────────────
    add_heading(doc, "Executive Summary", level=1)
    add_body(doc,
             "MedGuard AI is a working prototype that demonstrates how three established AI patterns — "
             "natural-language search, explainable risk scoring, and document-grounded retrieval — can be "
             "combined into a single decision-support cockpit for any large claim portfolio.")
    add_body(doc,
             "We have built and tested the full architecture end-to-end on synthetic medical records, "
             "because medical text is the hardest case in insurance. The same modular pattern is directly "
             "re-applicable to The Hartford's two largest opportunities:")
    add_bullets(doc, [
        "Group Benefits — short-term-to-long-term disability triage and absence management",
        "Middle & Large Commercial — Workers' Compensation severity reduction and Life / Specialty underwriting acceleration",
    ])
    add_body(doc,
             "A 1 % improvement in any of these workflows is a multi-million-dollar annual outcome.",
             bold=True, color=NAVY)

    # ── 1 LINE OF BUSINESS ────────────────────────────────
    add_heading(doc, "1.  Our Line of Business", level=1)
    add_body(doc,
             "The Hartford operates across three lines that all share one structural challenge: the "
             "most valuable signal lives in unstructured medical text that today is read manually.")
    add_table(doc,
              headers=["Line of Business", "What we cover", "Why this matters here"],
              rows=[
                  ["Middle & Large Commercial",
                   "Property, Liability, Workers' Comp for mid-market employers ($25M – $1B revenue) and Fortune-1000 risk programs",
                   "Medical evidence drives Workers' Comp severity; opioid and polypharmacy red flags are the largest controllable cost driver"],
                  ["Group Benefits",
                   "Disability (STD / LTD), Absence, FMLA, Group Life — #2 U.S. carrier in group disability",
                   "Nurse case managers read clinical evidence sequentially; earlier risk identification has direct dollar impact"],
                  ["Specialty & Life",
                   "Underwriting, specialty risk",
                   "Underwriters spend 30 – 60 minutes per APS letter; semantic retrieval reduces cycle time meaningfully"],
              ],
              first_col_bold=True)
    add_body(doc, "", size=4)
    add_body(doc, "Common thread:  every meaningful loss decision involves clinical evidence — "
                  "and that evidence is unstructured.", bold=True, color=NAVY)

    # ── 2 OPPORTUNITY ─────────────────────────────────────
    add_heading(doc, "2.  The Opportunity", level=1)
    add_callout_box(doc,
        "INDUSTRY BENCHMARK",
        "$15 – 30 M  ·  annual value of a 1 % reduction in STD-to-LTD conversion "
        "for a top-3 U.S. group disability carrier")
    add_body(doc, "Today's reality", bold=True, color=NAVY, size=12)
    add_bullets(doc, [
        "Nurse case managers read records sequentially",
        "Drug-interaction red flags surface only after escalation",
        "Underwriters spend 30 – 60 minutes per Attending Physician Statement",
        "Multilingual members receive English-only service",
        "Risk indicators sit in unstructured text the system cannot query",
    ])
    add_body(doc,
             "MedGuard AI gives every reviewer an AI co-pilot — risk surfaces in seconds, not days.",
             bold=True, color=HARTFORD_RED, size=12)

    # ── 3 ARCHITECTURE ────────────────────────────────────
    add_heading(doc, "3.  Architecture", level=1)
    add_body(doc, "Four modules · one pipeline · zero proprietary APIs.",
             bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    flow = ("📄 Data   ▶   🤖 JARVIS   ▶   📊 NotebookLM   ▶   "
            "🛡️ GARMA   ▶   📈 Cockpit")
    add_run(p, flow, size=12, bold=True, color=NAVY,
            font="Consolas")

    add_table(doc,
              headers=["Module", "Inspired by", "What it does"],
              rows=[
                  ["📄  Data Engine",  "FHIR / HL7",
                   "Generates 250 synthetic patients, validated with Pydantic — production swap point for Guidewire"],
                  ["🤖  JARVIS",       "AI-Augmented Engineering",
                   "Plain-English clinical querying with sub-second response, no API key required"],
                  ["🛡️  GARMA",       "GenAI Risk Mitigation Advisor",
                   "Composite 6-factor risk score, drug-interaction detection, anomaly engine, recommendations"],
                  ["📊  NotebookLM",   "Google NotebookLM",
                   "Vector retrieval over clinical text using FAISS + sentence-transformers"],
                  ["📈  Cockpit",      "Streamlit",
                   "KPI dashboard, alerts, clinician reports, multilingual voice in 19 languages"],
              ],
              first_col_bold=True)
    add_body(doc, "")
    add_body(doc, "Every component is open-source.  No vendor lock-in.  No API key required for the demo.",
             color=GRAY, size=10)

    # ── 4 WHAT WE BUILT ──────────────────────────────────
    add_heading(doc, "4.  What We Have Built", level=1)
    add_table(doc,
              headers=["Component", "Status", "Notes"],
              rows=[
                  ["Synthetic patient generator (250 records)", "✓ Complete",
                   "FHIR-style schema with intentional risk patterns embedded"],
                  ["JARVIS NL agent", "✓ Complete",
                   "Clinical taxonomy, scored relevance, 8 preset queries"],
                  ["GARMA risk engine", "✓ Complete",
                   "6-factor weighted composite, drug-interaction matrix, anomaly rules"],
                  ["NotebookLM RAG layer", "✓ Complete",
                   "Sentence-transformer embeddings, FAISS index"],
                  ["Streamlit cockpit", "✓ Complete",
                   "5 pages, KPI cards, critical-alert feed, radar charts"],
                  ["Multilingual voice", "✓ Complete",
                   "19 languages, online + offline (Windows SAPI) modes"],
                  ["Corporate-network resilience", "✓ Complete",
                   "Handles Hartford / Zscaler SSL inspection via truststore"],
                  ["Polished pptx deck and Word document", "✓ Complete",
                   "MedGuard_AI.pptx and MedGuard_AI.docx"],
              ],
              first_col_bold=True)

    # ── 4b WHO USES IT ───────────────────────────────────
    add_heading(doc, "4b. Who Uses MedGuard AI — Personas", level=1)
    add_body(doc, "Four Hartford personas that all read medical text today.",
             color=GRAY)
    add_table(doc,
              headers=["Persona", "Line of Business", "Today", "With MedGuard AI"],
              rows=[
                  ["🩺  Nurse Case Manager",     "Group Benefits — LTD",
                   "Reads 40-page records sequentially",
                   "Ranked work queue + auto-flagged drug interactions"],
                  ["💼  Claim Handler",           "Workers' Compensation",
                   "No clinical training, must set reserves",
                   "Plain-English risk score with explanation"],
                  ["📋  Underwriter",             "Life / Specialty",
                   "30 – 60 minutes per APS letter, manual",
                   "Asks APS questions, gets cited paragraph answers"],
                  ["📞  Member-Services Agent",  "Group Benefits / WC",
                   "English-only callbacks, escalates to interpreter",
                   "Instant voice reply in 19 languages"],
              ],
              first_col_bold=True)

    # ── 4c DAY IN THE LIFE ──────────────────────────────
    add_heading(doc, "4c. A Day in the Life — Priya, LTD Nurse Case Manager", level=1)
    add_body(doc, "Monday morning, 8:00 a.m. — total elapsed time: 5 minutes.",
             color=GRAY)
    add_table(doc,
              headers=["Time", "Action", "What happens", "Why it matters"],
              rows=[
                  ["8:00", "Opens the cockpit",
                   "250 claimants, 47 critical alerts, 18 drug interactions caught.",
                   "5 seconds vs 45 minutes in Excel today"],
                  ["8:01", "Asks JARVIS in plain English",
                   "“Find high-risk elderly patients on warfarin” → 10 ranked patients.",
                   "No SQL, no analyst needed"],
                  ["8:02", "Drills into the worst patient (GARMA)",
                   "Brian Yang, 71M, risk 91 % — radar shows polypharmacy + drug interaction.",
                   "Decision-grade explanation, not a black box"],
                  ["8:03", "Reads the recommendation",
                   "Warfarin + Aspirin → severe bleeding risk. Recommends INR check + PPI.",
                   "Picks up the phone and calls treating physician"],
                  ["8:05", "Documents and moves on",
                   "Downloads clinician report, attaches to claim, audit-trail recorded.",
                   "Next claim — repeatable, auditable workflow"],
              ],
              first_col_bold=True)

    # ── 4d THREE CUSTOMER USE CASES ─────────────────────
    add_heading(doc, "4d. Three Concrete Customer Use Cases", level=1)
    add_body(doc, "Same product, three lines of business.", color=GRAY)
    add_table(doc,
              headers=["Use Case", "Workflow", "Sample Query / Action", "Result"],
              rows=[
                  ["Group Benefits — STD-to-LTD triage",
                   "Spot the 5 % of STD claims that will convert to LTD on day 1",
                   "“Show STD claims week-3 with opioids and no return-to-work plan.”",
                   "Top 3 get nurse intervention → 1 % conversion reduction = $15 – 30 M / yr"],
                  ["Workers' Comp — catastrophic-claim watch",
                   "Detect opioid + benzo combinations driving claim escalation",
                   "“Which open WC claims have opioid + benzo prescriptions?”",
                   "9 claims escalated to medical director → ~25 % shorter duration"],
                  ["Life / Specialty UW — APS acceleration",
                   "Eliminate the 30 – 60 minutes of manual APS reading per case",
                   "“Any history of cardiac events in this APS?”",
                   "Cited paragraph returned → bind decision in 5 min, throughput +25 %"],
              ],
              first_col_bold=True)


    # ── 5 DEMO FLOW ───────────────────────────────────────
    add_heading(doc, "5.  Live Demonstration Flow (5 minutes)", level=1)
    add_table(doc,
              headers=["Time", "Section", "What happens on screen"],
              rows=[
                  ["0:00 – 0:20", "Title",                "Set the stage"],
                  ["0:20 – 0:50", "Line of business",     "Hartford context"],
                  ["0:50 – 1:10", "Opportunity",          "$15 – 30M anchor"],
                  ["1:10 – 2:00", "Architecture",         "Four-module pipeline"],
                  ["2:00 – 2:40", "DEMO 1 — Cockpit",     "KPIs, critical alerts, risk distribution"],
                  ["2:40 – 3:30", "DEMO 2 — JARVIS",      "NL query → 10 patients → multilingual voice"],
                  ["3:30 – 4:10", "DEMO 3 — GARMA",       "Radar chart, drug interaction, recommendation"],
                  ["4:10 – 4:40", "Hartford mapping",     "LTD / WC / Underwriting"],
                  ["4:40 – 4:55", "Path to production",   "Azure, governance, compliance"],
                  ["4:55 – 5:00", "Close + Q&A",          "“A pattern, not a prototype”"],
              ],
              first_col_bold=True)

    # ── 6 HARTFORD APPLICATIONS ──────────────────────────
    add_heading(doc, "6.  Hartford Applications", level=1)
    add_body(doc, "Same architecture, three immediate Hartford use cases.",
             bold=True, color=NAVY)
    add_table(doc,
              headers=["MedGuard module", "Group Benefits LTD", "Workers' Comp", "Life / Specialty UW"],
              rows=[
                  ["🤖  JARVIS",
                   "Show LTD claimants with no RTW plan",
                   "Find catastrophic claims on opioids",
                   "Underwriter NL search over APS"],
                  ["🛡️  GARMA",
                   "STD → LTD conversion risk score",
                   "Opioid / polypharmacy red flags",
                   "Mortality composite scoring"],
                  ["📊  NotebookLM",
                   "RAG over IME and APS reports",
                   "Treatment-note semantic search",
                   "MIB / APS PDF retrieval"],
                  ["📈  Cockpit",
                   "Nurse case-manager triage queue",
                   "Claim-handler severity dashboard",
                   "Underwriter decision support"],
              ],
              first_col_bold=True)

    add_body(doc, "")
    add_body(doc, "Estimated value (industry benchmarks)", bold=True, color=NAVY, size=12)
    add_table(doc,
              headers=["Lever", "Impact"],
              rows=[
                  ["1 % LTD-duration reduction",      "$15 – 30 M annually"],
                  ["Underwriter cycle time",           "~25 % faster APS review"],
                  ["Multilingual self-service",        "5 – 10 % call deflection"],
              ],
              first_col_bold=True)

    # ── 7 PATH TO PRODUCTION ─────────────────────────────
    add_heading(doc, "7.  Path to Production", level=1)
    add_table(doc,
              headers=["#", "Workstream", "What changes"],
              rows=[
                  ["01", "Data sources",         "Replace synthetic generator with Guidewire ClaimCenter, APS feeds, Rx data"],
                  ["02", "Enterprise AI",         "Azure OpenAI + Azure Translator / Speech (existing Hartford licenses)"],
                  ["03", "Security & compliance", "Hartford SSO · PHI guardrails · audit logging · HITRUST · HIPAA"],
                  ["04", "Hosting",               "Azure App Service / AKS behind Zscaler · zero data egress"],
                  ["05", "Model governance",      "Register GARMA in Hartford AI Model Risk inventory · NAIC-aligned monitoring"],
              ],
              first_col_bold=True)

    # ── 8 Q&A ─────────────────────────────────────────────
    add_heading(doc, "8.  Q & A — Anticipated Questions", level=1)
    add_table(doc,
              headers=["Question", "One-line answer"],
              rows=[
                  ["Is the data real?",
                   "Synthetic — Faker against FHIR-style Pydantic schemas. No PHI ever touches the demo."],
                  ["What LLM is used?",
                   "None in the demo, intentionally — JARVIS uses keyword + clinical-taxonomy matching. Azure OpenAI slots in for production."],
                  ["How is the risk score validated?",
                   "Today, transparent weighted composite of six clinical factors for explainability. Production validation against historical LTD outcomes plus AI inventory registration."],
                  ["Why Streamlit?",
                   "Speed of iteration. Cockpit is a front-end; modules underneath are framework-agnostic Python."],
                  ["Privacy / HIPAA?",
                   "Demo is synthetic. Production: Azure-private LLMs, no data egress, full audit trail, PHI tokenization at the edge."],
                  ["Time to a real Hartford pilot?",
                   "With committed data access and Azure provisioning: 8 – 12 weeks to a controlled LTD pilot on a single book of business."],
              ],
              first_col_bold=True)

    # ── 9 CLOSING ─────────────────────────────────────────
    add_heading(doc, "9.  Closing", level=1)
    add_callout_box(doc,
        "MEDGUARD AI",
        "A pattern, not just a prototype.")
    add_bullets(doc, [
        "Four modules · one architecture · zero vendor lock-in",
        "End-to-end working demo in five minutes",
        "Re-applicable across Disability, Workers' Compensation, and Underwriting",
        "Clear path to Azure-hosted, HIPAA-compliant production",
    ])
    add_body(doc, "")
    add_body(doc, "Thank you.", bold=True, color=NAVY, size=14)

    # ── Footer line ───────────────────────────────────────
    add_horizontal_rule(doc, color="0A2540", size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Confidential — for internal demo  ·  The Hartford  ·  2026",
            size=9, color=GRAY)

    out = Path(__file__).parent / "MedGuard_AI.docx"
    doc.save(str(out))
    print(f"✓ Document written to {out}")


if __name__ == "__main__":
    build()
